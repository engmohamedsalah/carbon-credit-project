import docx
import os

def extract_text_from_docx(file_path):
    """Extract text from a docx file."""
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text.append(paragraph.text)
    
    return '\n'.join(full_text)

def main():
    # Define file paths
    literature_review_path = "/home/ubuntu/upload/Literature review - AI Carbon Credit Verification Opportunities Limitations and Ethical Considerations.docx"
    proposal_path = "/home/ubuntu/upload/Proposal_ Carbon Credit Verification SaaS Application.docx"
    
    # Extract and save literature review content
    literature_review_text = extract_text_from_docx(literature_review_path)
    with open("/home/ubuntu/carbon_credit_project/literature_review.txt", "w") as f:
        f.write(literature_review_text)
    
    # Extract and save proposal content
    proposal_text = extract_text_from_docx(proposal_path)
    with open("/home/ubuntu/carbon_credit_project/proposal.txt", "w") as f:
        f.write(proposal_text)
    
    print("Extraction complete. Files saved to:")
    print("- /home/ubuntu/carbon_credit_project/literature_review.txt")
    print("- /home/ubuntu/carbon_credit_project/proposal.txt")

if __name__ == "__main__":
    main()
